from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = "docs/Bayangan_yang_Tertinggal_Suggested_Changes_15_Minute_Playthrough.docx"


COLORS = {
    "ink": RGBColor(31, 35, 40),
    "muted": RGBColor(90, 96, 105),
    "red": RGBColor(132, 32, 41),
    "cream": "F7F2EA",
    "dark": "211F1C",
    "line": "D7CFC3",
    "soft_red": "F2DEDD",
    "soft_green": "E2EFE6",
    "soft_gold": "F5E9C8",
}


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D7CFC3", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + margin))
        if node is None:
            node = OxmlElement("w:" + margin)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def style_paragraph(paragraph, size=10.5, color=None, bold=False, italic=False, before=0, after=6, line=1.08):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    for run in paragraph.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = "Heading {}".format(level)
    run = paragraph.add_run(text)
    run.font.name = "Aptos Display"
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = COLORS["red"]
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(6)
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = COLORS["ink"]
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_body(doc, text, after=7):
    paragraph = doc.add_paragraph(text)
    style_paragraph(paragraph, size=10.5, color=COLORS["ink"], after=after)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.add_run("- " + text)
    style_paragraph(paragraph, size=10.3, color=COLORS["ink"], after=3)
    return paragraph


def add_callout(doc, title, body, fill):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, color=fill, size="4")
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    p_title = cell.paragraphs[0]
    p_title.add_run(title).bold = True
    style_paragraph(p_title, size=11.2, color=COLORS["red"], after=4)
    p_body = cell.add_paragraph(body)
    style_paragraph(p_body, size=10.2, color=COLORS["ink"], after=0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths_cm)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        cell.text = header
        shade_cell(cell, COLORS["dark"])
        set_cell_border(cell, color=COLORS["dark"])
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(paragraph, size=9.4, color=RGBColor(255, 255, 255), bold=True, after=0)
    for row_data in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cell = row_cells[idx]
            cell.text = str(value)
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 1 and len(headers) <= 4 else WD_ALIGN_PARAGRAPH.LEFT
                style_paragraph(paragraph, size=9.4, color=COLORS["ink"], after=0, line=1.05)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = COLORS["ink"]

    for style_name in ("List Bullet", "List Number"):
        styles[style_name].font.name = "Aptos"
        styles[style_name].font.size = Pt(10.3)


def build_doc():
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Bayangan yang Tertinggal")
    run.font.name = "Aptos Display"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = COLORS["red"]

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Suggested Changes and Full Storyline for a 15-Minute Interactive Playthrough")
    run.font.name = "Aptos"
    run.font.size = Pt(13)
    run.font.color.rgb = COLORS["muted"]

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared for project review and sharing").italic = True
    style_paragraph(meta, size=9.5, color=COLORS["muted"], after=14)

    add_callout(
        doc,
        "Recommendation",
        "Adapt the prototype into a focused 15-minute adventure-horror visual novel. The player explores Batu Layar, collects meaningful tools and clues, survives short panic moments, and determines whether the pocong is destroyed, escaped, or properly released.",
        COLORS["cream"],
    )

    add_heading(doc, "1. Current Issue", 1)
    add_body(
        doc,
        "The current build already has meaningful story branches, especially in Beats 2-4. However, much of the experience is still passive because the player mostly reads dialogue, selects standard menu choices, and waits through long pacing beats. The story can become stronger if it is shorter, more focused, and more visibly responsive to player decisions.",
    )
    add_body(
        doc,
        "The present script is approximately 6,500+ words before pauses, menu time, and player hesitation. For a 15-minute playthrough, the safer target is about 2,500-3,200 words. Horror also needs breathing room, so the reduction should prioritize repeated explanation, duplicate choice gates, and long aftermath passages.",
    )

    add_heading(doc, "2. Target Experience", 1)
    for item in [
        "A complete playthrough should take about 13-16 minutes.",
        "The player should understand the main story without needing all optional details.",
        "Every major interaction should either reveal a clue, change MC's condition, or affect the ending.",
        "The pocong should remain scary, but the emotional truth should be discoverable.",
        "The best ending should feel earned through investigation, not simply selected at the end.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. Proposed 15-Minute Structure", 1)
    add_table(
        doc,
        ["Section", "Target Time", "Purpose", "Suggested Change"],
        [
            ["Cold open: Nayan", "1.5 min", "Establish threat", "Keep the death scene but shorten repeated buildup and pauses."],
            ["Briefing and Melur call", "2 min", "Introduce case", "Merge restaurant, phone call, and folklore briefing into one compact scene."],
            ["Arrival and first attack", "3 min", "Create danger", "Keep Beat 2 but reduce to two high-pressure decisions."],
            ["Investigation hub", "4 min", "Let player uncover truth", "Replace linear Beat 3 with three interactive investigation locations."],
            ["Final confrontation", "4 min", "Resolve story", "Compress Beat 4 into three rounds with clue-gated options."],
            ["Ending", "0.5 min", "Show consequence", "Deliver short ending based on how much truth the player found."],
        ],
        [4.0, 2.2, 4.0, 7.0],
    )

    add_heading(doc, "4. Story Trimming Plan", 1)
    add_table(
        doc,
        ["Area", "Keep", "Trim or Merge"],
        [
            ["Beat 1 cold open", "Nayan, silence, pocong approach, death.", "Reduce repeated 'Duk' lines and long pauses."],
            ["Restaurant scenes", "Hafiz mention, Melur's request, key pocong rules.", "Merge three scenes into one briefing. Remove duplicate folklore explanations."],
            ["Beat 2 attack", "First real player danger, injury consequences.", "Reduce to approach choice and shriek reaction choice."],
            ["Beat 3 investigation", "Mak Ros, burial ground, old house, mother truth.", "Make these selectable locations instead of a long linear route."],
            ["Beat 4 finale", "Knowledge vs ignorance ending logic.", "Cut five combat rounds to three decisive rounds."],
            ["Abandon ending", "Moral consequence of leaving.", "Keep only if used as one major late-game choice."],
        ],
        [4.2, 6.2, 6.8],
    )

    add_heading(doc, "5. Interaction Additions", 1)
    add_table(
        doc,
        ["Feature", "Player Experience", "Implementation Note"],
        [
            ["Case Notes", "Player sees clues collected during play.", "Use existing variables such as learned_rushed_burial, found_burial_site, and mother_told_truth."],
            ["Investigation Hub", "Player chooses where to investigate first.", "Create a menu or screen with Mak Ros, burial ground, old house, and prepare for night."],
            ["Clickable Inspection", "Player clicks objects at key locations.", "Use Ren'Py screens for grave soil, torn cloth, old photo, or prayer object."],
            ["Timed Choices", "Fight scenes feel urgent and tense.", "Use timed menu screens for cover ears, dodge, recite, or stand still."],
            ["Choice Feedback", "Player understands consequences.", "Show small notifications: clue added, injury worsened, opening gained."],
            ["Clue-Gated Ending", "Best ending feels earned.", "Only unlock 'speak/release' route if enough truth has been discovered."],
        ],
        [3.4, 6.5, 7.3],
    )

    add_heading(doc, "6. Revised Story Flow", 1)
    add_body(doc, "The recommended playable route should follow this concise spine:")
    for item in [
        "Nayan is killed in Batu Layar, establishing the pocong as a serious threat.",
        "MC hears about the village, receives Melur's call, and learns the minimum folklore needed: the kafan knots matter, rushed burial matters, and release is different from killing.",
        "MC arrives and survives the first encounter. The player's reactions determine injury, fear, and whether MC notices a useful movement pattern.",
        "During the day, the player investigates three important leads: witness, grave, and old house. Each lead adds a clue.",
        "At night, the pocong attacks again. The player's previous clues and injuries affect available choices.",
        "If the player learned the truth, MC can release the pocong properly. If not, MC can only defeat it or fail.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. Full New Version Storyline", 1)
    add_body(
        doc,
        "This version keeps the original main story: MC is called to Batu Layar, discovers the pocong is Melur's dead brother rather than a simple monster, learns he was trapped by a rushed burial and a hidden sacrifice, then decides whether to release or destroy him. The difference is that the player now uncovers that truth through tools, inspection, and short adventure-game objectives.",
    )
    add_table(
        doc,
        ["Chapter", "Time", "Location", "Story Purpose"],
        [
            ["Prologue", "0:00-1:15", "Village path", "Show Nayan's death and establish the pocong's sound, movement, and danger."],
            ["Chapter 1", "1:15-3:00", "Restoran Zulkifli", "Brief MC, introduce Melur, and let player choose what question to prioritize."],
            ["Chapter 2", "3:00-5:00", "Batu Layar entrance", "First playable encounter; player survives and learns the pocong reacts to behavior."],
            ["Chapter 3", "5:00-10:30", "Investigation hub", "Player explores three locations, collects tools, and builds understanding."],
            ["Chapter 4", "10:30-13:30", "Night preparation and attack", "Collected tools become usable under pressure."],
            ["Chapter 5", "13:30-15:00", "Burial ground finale", "Ending is decided by truth, tools, and player intent."],
        ],
        [2.6, 2.0, 4.0, 8.2],
    )

    add_heading(doc, "Prologue: The Last Walk Home", 2)
    add_body(
        doc,
        "Nayan walks home alone through Batu Layar at 2:47 a.m. The scene is short and mostly linear, but the player gets one small interaction: look back, call out, or keep walking. No choice saves him. The purpose is to teach the player the pocong's language: silence, distant crying, then the heavy repeated hop.",
    )
    add_body(
        doc,
        "The attack ends in a flash and scream. The player never sees the pocong clearly, only white cloth at the edge of the road. This keeps the monster frightening while setting up later investigation clues: the sound pattern, the tied body movement, and the bleeding ears.",
    )

    add_heading(doc, "Chapter 1: The Call and the Kit", 2)
    add_body(
        doc,
        "In Restoran Zulkifli, Hafiz mentions the Batu Layar deaths. Melur calls shortly after and asks MC to go there. Her wording stays emotionally important: she does not ask MC to kill the pocong; she asks him to help her brother be released properly.",
    )
    add_body(
        doc,
        "Abang Zul gives MC a compact field kit. The player sees the first tool choice, but this should be simple rather than a large inventory puzzle. Suggested starting items: flashlight, phone camera, keris/small blade, and empty case notes. Optional protection items such as salt or tasbih must be found in Batu Layar, not given immediately.",
    )
    add_body(
        doc,
        "Player interaction: choose two questions before leaving. Asking about rushed burial increases understanding. Asking what pocong wants unlocks empathy language in the finale. Asking how to fight gives a combat option but does not help the best ending by itself.",
    )

    add_heading(doc, "Chapter 2: First Night in Batu Layar", 2)
    add_body(
        doc,
        "MC arrives as a villager is being chased. This replaces a long action scene with one high-pressure survival sequence. The player can warn the villager, stand still, grab a branch, or run. The best clue comes from standing still or observing: the pocong slows for a moment before attacking again.",
    )
    add_body(
        doc,
        "A timed shriek choice follows. Covering ears prevents mental shock but may cause physical injury. Standing firm preserves observation but raises fear. Using the phone camera during the blackout captures a blurred image of tied kafan cloth, adding an early visual clue. The encounter ends with MC alive, injured or shaken, and aware that brute force is not enough.",
    )

    add_heading(doc, "Chapter 3: Day Investigation Hub", 2)
    add_body(
        doc,
        "The player gets three investigation stops before nightfall. They may choose the order, but each location takes about 90 seconds. After all three are visited, the game automatically moves to final preparation. This keeps the adventure feeling open without breaking the 15-minute cap.",
    )
    add_table(
        doc,
        ["Location", "Tool Interaction", "Clue Gained", "Consequence"],
        [
            ["Mak Ros's house", "Use phone camera at the window or ask about the sound.", "The pocong repeats the same route and reacts to people who run.", "Unlocks safer dodge/stand-still options later."],
            ["Burial ground", "Use flashlight to inspect soil, cloth, and knot marks.", "The grave was rushed and the kafan knot was never opened.", "Unlocks ritual objective: open the knot."],
            ["Old family house", "Find old letter, tasbih, and mother's testimony.", "The pocong was Melur's brother, blamed after stopping a dark bomoh.", "Unlocks name/empathy options in the finale."],
        ],
        [3.6, 4.8, 5.4, 4.4],
    )
    add_body(
        doc,
        "The old family house is the emotional center. The mother does not give a long monologue unless the player has found at least one earlier clue. If the player arrives unprepared, she only gives partial truth. If the player brings the kafan thread or photo evidence, she tells the fuller story: her son made a terrible sacrifice to stop a greater evil, was misunderstood, buried in anger, and left trapped.",
    )

    add_heading(doc, "Chapter 4: Preparing for Night", 2)
    add_body(
        doc,
        "As the sun falls, the player reviews case notes and chooses how to prepare. This is the adventure-game payoff: tools become intentions. The keris can be held as a weapon or kept sheathed. The tasbih can be used for prayer. The flashlight can mark the grave path. The phone camera can review ghost images. The old letter can reveal the correct name.",
    )
    add_body(
        doc,
        "A short scare interrupts preparation. The pocong appears at the village edge. Timed options depend on tools: shine flashlight, recite with tasbih, throw salt if found, raise keris, or stay still. Aggressive choices help survival but increase fear or anger. Respectful choices require clues but move the player closer to release.",
    )

    add_heading(doc, "Chapter 5: The Burial Ground Finale", 2)
    add_body(
        doc,
        "The finale happens at the grave, not a random battlefield. The pocong blocks the path while the cloth around him shakes as if pulled by invisible knots. The player must decide what the whole investigation means: is this a monster to stop, a victim to release, or a danger too strong to face?",
    )
    add_body(
        doc,
        "Final sequence: first survive the shriek, then approach or attack, then choose the finishing action. The best route requires the old letter or mother's truth, the burial clue, and a ritual item such as tasbih. The player calls him by name, explains that his mother remembers him, opens the knot, and recites the prayer. The spirit speaks once, asks MC to watch over his son, then leaves.",
    )
    add_body(
        doc,
        "If the player lacks truth, the same tools produce a different result. The keris can force the pocong down, salt can hold him back, and prayer can silence him temporarily, but the narration makes clear that this is not release. He is stopped, not understood.",
    )

    add_heading(doc, "8. Chapter Loading Screens", 1)
    add_body(
        doc,
        "Each chapter should begin with a short static loading or chapter-introduction screen. This does not need animation. A still background, chapter title, location, time, and character lineup is enough. The purpose is to orient the player quickly, reset pacing between scenes, and make the 15-minute structure feel intentional.",
    )
    add_body(
        doc,
        "Recommended format: dark scenic background or blurred location art, chapter number, short subtitle, involved characters, and one atmospheric sentence. Keep each screen on display for about 2-3 seconds, or allow the player to click to continue. These cards can also hide small loading delays if heavier assets are added later.",
    )
    add_table(
        doc,
        ["Chapter", "Static Card Title", "Characters Introduced", "Mood / Purpose"],
        [
            ["Prologue", "Pukul 2:47 Pagi", "Nayan, unseen pocong", "Teaches the threat through silence, crying, and the sound of hopping."],
            ["Chapter 1", "Restoran Zulkifli", "MC, Hafiz, Melur, Abang Zul", "Frames the case and introduces the difference between killing and releasing."],
            ["Chapter 2", "Malam Pertama di Batu Layar", "MC, villagers, pocong", "Signals that the player is now inside danger, not just hearing about it."],
            ["Chapter 3", "Sisa Yang Tertinggal", "MC, Mak Ros, mother, son", "Introduces investigation mode and the people connected to the old tragedy."],
            ["Chapter 4", "Menunggu Malam", "MC, pocong", "Shows preparation, tool review, and the feeling that time is running out."],
            ["Chapter 5", "Simpulan Terakhir", "MC, arwah/pocong", "Frames the finale as release versus destruction."],
        ],
        [2.4, 4.5, 5.4, 5.8],
    )
    add_body(
        doc,
        "Character introduction should be functional, not overly explanatory. Example: 'MC - penyiasat kes ghaib', 'Melur - adik kepada arwah', 'Mak Ros - saksi malam pertama', 'Arwah - roh yang belum dilepaskan'. This helps players remember who matters without adding long exposition inside dialogue.",
    )
    add_body(
        doc,
        "Visual direction: use one consistent template for all chapter cards. Suggested composition: large chapter title at center-left, small character list at bottom-left, location/time at top-right, and a faint silhouette or location image in the background. Avoid busy UI; the card should feel like a calm breath before the next scare.",
    )

    add_heading(doc, "9. Adventure Tool and Clue Logic", 1)
    add_table(
        doc,
        ["Tool / Clue", "How Player Gets It", "Final Use"],
        [
            ["Flashlight", "Starter kit.", "Reveals knot marks and helps navigate the grave path."],
            ["Phone camera", "Starter kit.", "Captures spirit traces and confirms the pocong is repeating old memory."],
            ["Keris / small blade", "Starter kit.", "Can fight, but best use is carefully opening the kafan knot."],
            ["Kafan thread", "Burial ground inspection.", "Proof that the burial was rushed and the knot remains unresolved."],
            ["Old letter", "Mother's house.", "Reveals identity, sacrifice, and the name needed for release."],
            ["Tasbih", "Mother's house or prayer corner.", "Stabilizes the release ritual and unlocks prayer option."],
            ["Salt", "Optional kitchen / village house pickup.", "Blocks one attack but cannot solve the haunting."],
        ],
        [3.3, 5.2, 8.3],
    )

    add_heading(doc, "10. Ending Logic", 1)
    add_table(
        doc,
        ["Ending", "Requirement", "Outcome"],
        [
            ["Release Ending", "Mother truth + burial clue + enough investigation.", "MC speaks to the arwah, opens the knot, and lets him go properly."],
            ["Ignorance Ending", "Player survives but lacks key truth.", "MC destroys or suppresses the pocong, but the emotional truth remains unresolved."],
            ["Death Ending", "Too much damage or failed timed reactions.", "MC cannot survive the final encounter."],
            ["Abandon Ending", "Player chooses to leave before final preparation.", "The village is lost, and MC carries the consequence."],
        ],
        [3.5, 6.3, 7.4],
    )

    add_heading(doc, "11. Implementation Roadmap", 1)
    add_table(
        doc,
        ["Step", "Task", "Expected Result"],
        [
            ["1", "Rewrite the script outline into the new 15-minute structure.", "Clear scope before coding interaction systems."],
            ["2", "Merge Beat 1 restaurant, call, and folklore into one scene.", "Faster start and less exposition."],
            ["3", "Compress Beat 2 into two major pressure decisions.", "Stronger first encounter with clearer consequences."],
            ["4", "Rebuild Beat 3 as an investigation hub.", "Player actively chooses what to inspect and who to question."],
            ["5", "Add Case Notes and clue notifications.", "Player sees progress and understands discoveries."],
            ["6", "Compress Beat 4 and add clue-gated final options.", "Finale becomes shorter, more interactive, and more replayable."],
            ["7", "Playtest and time the build.", "Adjust text, pauses, and menu count until the average run is 13-16 minutes."],
        ],
        [1.3, 8.3, 7.6],
    )

    add_heading(doc, "12. Success Criteria", 1)
    for item in [
        "Average first playthrough is 13-16 minutes.",
        "Players can explain who the pocong is and why he remains trapped.",
        "Players report that choices feel meaningful, not decorative.",
        "At least three interactions are active investigations rather than only dialogue menus.",
        "The final resolution changes based on clues found earlier.",
        "No scene feels like repeated exposition from an earlier scene.",
    ]:
        add_bullet(doc, item)

    add_callout(
        doc,
        "Priority Recommendation",
        "Start with the script trim first, then add Case Notes and the Beat 3 investigation hub. These two changes will create the biggest improvement in playtime, clarity, and player involvement without rebuilding the whole game.",
        COLORS["soft_green"],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Bayangan yang Tertinggal - Suggested Changes")
    style_paragraph(footer, size=8.5, color=COLORS["muted"], after=0)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
